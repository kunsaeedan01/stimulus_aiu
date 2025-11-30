import io
import os
from django.conf import settings
from django.utils import timezone
from django.core.files.base import ContentFile
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy


def _add_page_break_at_start(doc: Document):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)
    body = doc.element.body
    body.insert(0, paragraph._element)


def generate_application_docx(application):
    template_path = os.path.join(settings.BASE_DIR, "templates", "application_template.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at: {template_path}")

    papers = list(application.papers.all())
    if not papers:
        papers = [None]

    months_ru = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    today = timezone.now()
    today_str = f"{today.day:02d} {months_ru[today.month]} {today.year} г."

    owner_full_name = (
        application.owner.full_name or
        application.owner.get_full_name() or
        application.owner.email
    )

    rendered_docs = []

    for idx, paper in enumerate(papers):
        coauthors = list(paper.coauthors.all()) if paper else []
        signers = [{"full_name": owner_full_name}]
        for co in coauthors:
            if co.is_aiu_employee and co.full_name and co.full_name.strip():
                signers.append({"full_name": co.full_name.strip()})
        indexation = paper.indexation if paper else ""

        context = {
            "owner_full_name": owner_full_name,
            "owner_position": application.owner.position or "",
            "owner_subdivision": application.owner.subdivision or "",
            "owner_telephone": application.owner.telephone or "",
            "owner_email": application.owner.email or "",
            "today": today_str,

            "publication_date": (
                paper.publication_date.strftime("%d.%m.%Y")
                if paper and paper.publication_date else "___"
            ),
            "title": paper.title if paper else "",
            "journal": paper.journal_or_source if paper else "",
            "year": str(paper.year or "") if paper else "",
            "number": paper.number or "" if paper else "",
            "volume": str(paper.volume or "") if paper else "",
            "pages": paper.pages or "" if paper else "",
            "doi": paper.doi or "" if paper else "",
            "quartile": paper.quartile or "" if paper else "",
            "percentile": str(paper.percentile or "") if paper else "",
            "indexation": indexation,

            "coauthors": [
                {
                    "full_name": co.full_name or "",
                    "position": co.position or "",
                    "subdivision": co.subdivision or "",
                    "telephone": co.telephone or "",
                    "email": co.email or "",
                    "is_aiu": co.is_aiu_employee,
                }
                for co in coauthors
            ],
            "all_signatures": signers,
        }

        temp_tpl = DocxTemplate(template_path)
        temp_tpl.render(context)

        buf = io.BytesIO()
        temp_tpl.save(buf)
        buf.seek(0)

        doc = Document(buf)
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        if idx > 0:
            _add_page_break_at_start(doc)

        rendered_docs.append(doc)

    final_doc = rendered_docs[0]

    for src_doc in rendered_docs[1:]:
        for child in src_doc.element.body:
            final_doc.element.body.append(deepcopy(child))
        if src_doc.sections:
            src_sect = src_doc.sections[-1]
            dst_sect = final_doc.sections[-1]
            if dst_sect._sectPr is not None:
                dst_sect._sectPr.getparent().replace(
                    dst_sect._sectPr,
                    deepcopy(src_sect._sectPr)
                )
    output = io.BytesIO()
    final_doc.save(output)
    output.seek(0)

    filename = f"application_{application.id}.docx"
    return filename, ContentFile(output.read(), name=filename)